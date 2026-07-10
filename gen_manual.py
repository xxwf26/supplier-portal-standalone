"""生成供应商平台使用说明 DOCX"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 页面边距 ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── 字体默认值 ────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = '微软雅黑'
style.font.size = Pt(11)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def set_font(run, bold=False, size=11, color=None, name='微软雅黑'):
    run.font.name = name
    run.font.bold = bold
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.name = '微软雅黑'
    p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.runs[0].font.size = Pt(16)
    return p

def heading2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.name = '微软雅黑'
    p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.runs[0].font.size = Pt(13)
    return p

def heading3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.name = '微软雅黑'
    p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    p.runs[0].font.size = Pt(12)
    return p

def para(text, bold=False, size=11, color=None, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_font(run, bold=bold, size=size, color=color)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    set_font(run)
    return p

def note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run('📌 ' + text)
    set_font(run, color=(100, 100, 100), size=10)
    return p

def img_placeholder(caption=''):
    """图片占位框"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(f'[ 截图：{caption} ]' if caption else '[ 截图 ]')
    run.font.color.rgb = RGBColor(180, 180, 180)
    run.font.italic = True
    run.font.size = Pt(10)
    return p

def table(headers, rows, col_widths=None):
    """生成表格"""
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # 表头
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cell.paragraphs[0].add_run(h)
        set_font(run, bold=True, size=10)
        cell._element.get_or_add_tcPr().append(
            OxmlElement('w:shd') if False else _shd('#DDEEFF')
        )
        _set_cell_bg(cell, 'DDEEFF')
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=10)
    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()  # 表后空行
    return t

def _shd(color_hex):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    return shd

def _set_cell_bg(cell, hex_color):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

# ══════════════════════════════════════════════════════════
# 封面
# ══════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = t.add_run('供应商可视化平台')
set_font(run, bold=True, size=28)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = t2.add_run('功能使用说明')
set_font(run2, bold=True, size=20, color=(80, 80, 80))

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = t3.add_run('采购侧用户版')
set_font(run3, size=13, color=(120, 120, 120))

doc.add_page_break()

# ══════════════════════════════════════════════════════════
# 角色说明
# ══════════════════════════════════════════════════════════
heading1('角色说明')
img_placeholder('角色说明区域')
table(
    ['角色', '权限说明'],
    [
        ['管理员（admin）', '拥有全部操作权限，包括新建、编辑、删除、导入、导出、配置系统'],
        ['查看者（viewer）', '只能浏览画师信息，无法进行任何编辑或管理操作'],
    ],
    col_widths=[4, 11]
)
note('下文中标注「🔐 管理员」的功能仅管理员账号可见。')

# ══════════════════════════════════════════════════════════
# 一、登录
# ══════════════════════════════════════════════════════════
heading1('一、登录')
img_placeholder('登录页')
table(
    ['字段', '说明'],
    [
        ['用户名', '由管理员提供'],
        ['密码', '由管理员提供'],
        ['记住我', '勾选后登录状态保持 30 天，否则 8 小时后自动退出'],
    ],
    col_widths=[4, 11]
)
para('登录成功后直接进入主界面。')

# ══════════════════════════════════════════════════════════
# 二、主界面概览
# ══════════════════════════════════════════════════════════
heading1('二、主界面概览')
img_placeholder('主界面全局截图（标注各区域）')
para('PC 端布局分为三个区域：')
bullet('左侧固定面板：筛选条件')
bullet('顶部吸顶工具条：关键词搜索、排序、列数切换、管理操作按钮')
bullet('主内容区：画师卡片网格')
para('')
para('手机端：卡片竖向排列，筛选面板通过左下角浮动「筛选」按钮唤出。')

# ══════════════════════════════════════════════════════════
# 三、左侧筛选面板
# ══════════════════════════════════════════════════════════
heading1('三、左侧筛选面板')
para('位置：页面左侧固定栏（PC 端）；手机端点击左下角「筛选」按钮展开。')
para('筛选条件实时生效，多个条件同时选中时取交集（且的关系）。')

heading2('3.1 供应商类型')
img_placeholder('供应商类型筛选区域')
para('可选项：个人画师 / 艺术家 / 工作室 / 公司，支持多选。')

heading2('3.2 合作状态')
img_placeholder('合作状态筛选区域')
table(
    ['选项', '含义'],
    [
        ['库内合作', '已签约或正在合作中'],
        ['库外建联', '尚未正式合作，处于接触阶段'],
        ['已拉黑', '不再合作'],
        ['未填写', '状态未录入'],
    ],
    col_widths=[4, 11]
)
para('支持多选。')

heading2('3.3 擅长风格')
img_placeholder('擅长风格筛选区域')
para('标签形式多选。勾选「未填写」可筛选出尚未设置风格的画师。')

heading2('3.4 合作类型')
img_placeholder('合作类型筛选区域')
para('标签形式多选。勾选「未填写」可筛选出尚未设置合作类型的画师。')

heading2('3.5 历史参与项目')
img_placeholder('历史参与项目筛选区域')
para('选择具体项目名称，筛选曾参与该项目的画师，支持多选。')

heading2('3.6 报价范围')
img_placeholder('报价范围滑块区域')
para('拖动滑块设置最低价和最高价（单位：元），范围 0 ~ 10000。')
para('勾选「未填写报价」可额外包含未录入报价的画师。')
note('价格滑块和「未填写」可同时勾选，结果为两者并集。')

# ══════════════════════════════════════════════════════════
# 四、吸顶工具条
# ══════════════════════════════════════════════════════════
heading1('四、吸顶工具条')
img_placeholder('吸顶工具条全貌')
para('位置：主内容区顶部，滚动时固定在屏幕上方。')

heading2('4.1 关键词搜索')
img_placeholder('搜索输入框')
para('输入框支持搜索画师名称和备注内容，输入后实时过滤结果。')

heading2('4.2 排序')
img_placeholder('排序下拉菜单')
table(
    ['选项', '说明'],
    [
        ['默认排序', '按录入时间由新到旧'],
        ['评分从高到低', '未评分的排在末尾'],
        ['评分从低到高', '未评分的排在末尾'],
        ['合作频次从高到低', '—'],
        ['合作频次从低到高', '—'],
        ['最近更新', '按最后编辑时间排序'],
    ],
    col_widths=[5, 10]
)

heading2('4.3 列数切换（PC 端）')
img_placeholder('列数切换按钮')
para('右侧「自动 / 1 / 2 / 3 / 4 / 5 / 6」按钮，控制卡片每行显示数量。「自动」根据窗口宽度自适应。')

heading2('4.4 管理员操作按钮（🔐 管理员）')
img_placeholder('管理员操作按钮区域')
table(
    ['按钮', '功能'],
    [
        ['历史', '打开变更记录、导入批次、快照备份面板'],
        ['查重', '检测库内名称相似的画师'],
        ['导入 Excel', '批量导入画师数据'],
        ['新建供应商', '手动录入单个画师'],
    ],
    col_widths=[4, 11]
)

# ══════════════════════════════════════════════════════════
# 五、画师卡片
# ══════════════════════════════════════════════════════════
heading1('五、画师卡片')
img_placeholder('画师卡片示例（标注各元素）')
para('主内容区网格中，每张卡片代表一位画师。')

heading2('5.1 卡片展示的信息')
bullet('画师名称（大标题）')
bullet('类型标签：个人画师 / 艺术家 / 工作室 / 公司')
bullet('合作状态圆点：绿色=库内合作，蓝色=库外建联，灰色=已拉黑，橙色=未填写')
bullet('擅长风格标签')
bullet('合作类型标签')
bullet('报价：显示最低～最高价，或「未填写」')
bullet('评分：1 ~ 5 星')
bullet('合作频次：次数')
bullet('平台链接：小图标快速跳转')

heading2('5.2 选中卡片')
img_placeholder('卡片选中状态')
para('点击卡片左上角复选框（悬停时出现）选中该画师，可多选。')
para('选中后页面底部弹出浮动操作栏（详见第七节）。')

heading2('5.3 打开详情')
para('点击卡片主体区域（非复选框区域）打开该画师的详情弹窗。')

# ══════════════════════════════════════════════════════════
# 六、画师详情弹窗
# ══════════════════════════════════════════════════════════
heading1('六、画师详情弹窗')
para('触发方式：点击画师卡片。')

heading2('6.1 头部信息')
img_placeholder('详情弹窗头部')
para('显示画师名称、类型标签、合作状态标签，以及所属合作品类。')
para('管理员可在此区域点击「编辑」按钮进入编辑模式。')

heading2('6.2 概览')
img_placeholder('概览数据卡')
table(
    ['格', '内容'],
    [
        ['合作状态', '仅管理员可见'],
        ['合作频次', '与该画师的合作次数'],
        ['评分', '1 ~ 5 分星级'],
    ],
    col_widths=[4, 11]
)

heading2('6.3 作品展示')
img_placeholder('作品展示区域及灯箱')
para('横向滚动浏览画师作品图片，点击任意图片进入全屏灯箱查看大图。')
para('灯箱内支持以下操作：')
bullet('左右箭头或键盘方向键翻页')
bullet('底部缩略图快速跳转')
bullet('点击背景或右上角 ✕ 关闭')

heading2('6.4 擅长风格 & 合作类型')
img_placeholder('擅长风格和合作类型区域')
para('左右两栏并排显示，标签形式。')

heading2('6.5 报价参考')
img_placeholder('报价参考表格')
para('表格形式，每行为一种合作类型及其对应单价和计价单位。')

heading2('6.6 联系方式（🔐 管理员）')
img_placeholder('联系方式区域')
para('显示微信号、QQ 号、电话等结构化联系方式，仅管理员可见。')

heading2('6.7 平台链接')
img_placeholder('平台链接区域')
para('点击按钮在新标签页打开对应平台主页（小红书、微博、B 站等）。')

heading2('6.8 历史参与项目')
img_placeholder('历史参与项目区域')
para('展示画师曾参与的项目名称。')

heading2('6.9 备注与佐证图片')
img_placeholder('备注及佐证图片区域')
para('上方为文字备注内容，下方为佐证图片区（横向滚动，点击可全屏查看）。')

# ══════════════════════════════════════════════════════════
# 七、浮动选中操作栏
# ══════════════════════════════════════════════════════════
heading1('七、浮动选中操作栏（多选模式）')
img_placeholder('浮动操作栏全貌')
para('触发方式：勾选一个或多个画师卡片后，页面底部自动弹出。')
table(
    ['按钮', '功能'],
    [
        ['全选当前', '选中当前筛选/搜索结果的全部画师'],
        ['清除', '取消所有选中'],
        ['导出 PDF', '将选中画师的档案导出为 PDF 文件'],
        ['复制名单', '将选中画师的摘要信息复制到剪贴板'],
        ['批量删除', '🔐 管理员，删除选中的全部画师（可撤销）'],
    ],
    col_widths=[4, 11]
)
note('若当前有筛选或搜索条件，被隐藏的已选画师不会出现在导出/复制结果中，操作前会有提示。')

heading2('7.1 导出 PDF')
img_placeholder('导出 PDF 字段选择对话框')
para('点击「导出 PDF」后弹出字段选择对话框：')
bullet('默认全选所有字段')
bullet('可取消勾选不需要的字段（如联系方式、报价等敏感信息）')
bullet('「全选」/「全不选」快速切换')
bullet('选择偏好会记住，下次导出沿用上次设置')
para('')
para('确认后开始生成，每位画师占一页或多页，文件名格式为「画师档案_YYYY-MM-DD.pdf」。')

heading2('7.2 复制名单')
img_placeholder('复制名单示例')
para('复制为纯文本，格式示例：')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5)
run = p.add_run('1. 画师名称 | 类型: 个人画师 | 状态: 库内合作 | 风格: Q版、正比 | 报价: 角色原画 500元/张 | 评分: 4分 | 频次: 8')
run.font.name = 'Courier New'
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(60, 60, 60)

heading2('7.3 批量删除（🔐 管理员）')
img_placeholder('批量删除确认弹窗')
para('点击后弹出确认弹窗，列出将被删除的画师名单。确认后：')
bullet('成功提示中附带「撤销」按钮，10 秒内可点击一键恢复')
bullet('超过 10 秒后可在「历史」→「变更记录」中逐条撤回')

# ══════════════════════════════════════════════════════════
# 八、新建供应商
# ══════════════════════════════════════════════════════════
heading1('八、新建供应商（🔐 管理员）')
img_placeholder('新建供应商表单')
para('入口：顶部工具条右侧「新建供应商」按钮。')
para('')
table(
    ['字段', '必填', '说明'],
    [
        ['供应商名称', '✅', '输入时实时检测库内相似名称，避免重复录入'],
        ['供应商类型', '', '个人画师 / 艺术家 / 工作室 / 公司'],
        ['合作类型', '', '点击标签多选'],
        ['历史参与项目', '', '下拉选择'],
        ['报价参考', '', '可添加多条，每条包含合作类型、单价、计价单位'],
        ['联系方式', '', '可添加多条微信 / QQ / 电话'],
        ['擅长风格', '', '点击预设标签或输入自定义'],
        ['作品图片', '', '点击上传或粘贴图片'],
        ['平台链接', '', '选择平台后填入 URL'],
        ['备注', '', '文字 + 佐证图片'],
    ],
    col_widths=[4, 1.5, 9.5]
)
note('表单支持草稿自动保存。未提交直接关闭后，下次打开会提示恢复草稿。')
note('填写名称时若检测到相似画师，输入框下方显示警告名单，请确认是否重复后再提交。')

# ══════════════════════════════════════════════════════════
# 九、编辑画师信息
# ══════════════════════════════════════════════════════════
heading1('九、编辑画师信息（🔐 管理员）')
img_placeholder('编辑模式界面')
para('入口：打开画师详情弹窗 → 右上角「编辑」按钮。')
para('')
para('编辑模式下所有字段均可修改，与新建表单字段相同。特别说明：')
bullet('合作类型和擅长风格支持全部删空（保存后置为未填写）')
bullet('历史参与项目选择「未设置」可清空该字段')
bullet('编辑中途关闭弹窗时会弹出放弃确认')
bullet('编辑过程中自动保存草稿，意外退出后可恢复')
bullet('作品展示和备注佐证图片支持直接粘贴截图上传')

# ══════════════════════════════════════════════════════════
# 十、导入 Excel
# ══════════════════════════════════════════════════════════
heading1('十、导入 Excel（🔐 管理员）')
img_placeholder('Excel 导入弹窗')
para('入口：顶部工具条「导入 Excel」按钮。')
para('')
para('操作步骤：')
bullet('①  下载弹窗内提供的模板文件', level=0)
bullet('②  按模板格式填写画师数据', level=0)
bullet('③  上传 Excel 文件，预览数据确认无误', level=0)
bullet('④  点击「开始导入」', level=0)
para('')
para('导入结果：')
bullet('显示成功条数和失败条数，失败行列出具体原因')
bullet('导入完成后可在「历史」→「导入批次」查看，并支持整批撤销')
note('单次最多导入 500 条。')

# ══════════════════════════════════════════════════════════
# 十一、查重
# ══════════════════════════════════════════════════════════
heading1('十一、查重（🔐 管理员）')
img_placeholder('查重结果面板')
para('入口：顶部工具条「查重」按钮。')
para('')
para('系统自动检测库内名称相似的画师，分两类：')
table(
    ['类型', '判断标准'],
    [
        ['完全重复', '名称完全一致（不区分大小写）'],
        ['模糊相似', '名称中包含 2 个以上相同汉字片段（自动忽略公司后缀、地名等干扰词）'],
    ],
    col_widths=[4, 11]
)
para('结果以列表形式展示，每组显示涉及的画师名称，可直接点击跳转查看或删除。')

# ══════════════════════════════════════════════════════════
# 十二、历史与数据管理
# ══════════════════════════════════════════════════════════
heading1('十二、历史与数据管理（🔐 管理员）')
para('入口：顶部工具条「历史」按钮，打开数据管理面板。')
para('面板分三个标签页：')

heading2('12.1 变更记录')
img_placeholder('变更记录标签页')
para('按时间倒序显示所有数据操作（新增、编辑、删除、批量删除、撤销导入、撤回操作、恢复删除）。')
bullet('编辑类记录展开显示字段变更前后对比')
bullet('悬停单条记录显示「撤回」按钮，点击可还原该操作（支持新增、编辑、删除、批量删除）')
bullet('点击「刷新」获取最新记录')

heading2('12.2 导入批次')
img_placeholder('导入批次标签页')
para('列出所有通过 Excel 批量导入的历史批次，显示导入时间和条数。')
para('点击「撤销此批」可删除该批次导入的全部数据（需二次确认）。')

heading2('12.3 快照备份')
img_placeholder('快照备份标签页')
para('系统在每次批量导入和批量删除前自动创建数据库快照，最多保留 30 份。')
bullet('点击「立即备份」可手动创建快照')
bullet('点击「恢复」可将整个数据库回滚到该时刻的状态（不可逆，需二次确认）')

# ══════════════════════════════════════════════════════════
# 十三、系统配置
# ══════════════════════════════════════════════════════════
heading1('十三、系统配置（🔐 管理员）')
img_placeholder('系统配置页面全貌')
para('入口：页面右上角「系统配置」链接。')
para('')
table(
    ['分类', '说明'],
    [
        ['供应商类型', '固定字段，不支持增删'],
        ['合作类型', '可新增、编辑、排序、禁用'],
        ['擅长风格', '可新增、编辑颜色标签、排序、禁用'],
        ['合作状态', '固定逻辑，仅支持编辑显示名称'],
        ['历史参与项目', '可新增、编辑、排序、禁用'],
    ],
    col_widths=[4, 11]
)
note('在画师详情编辑中输入的自定义合作类型和风格标签，保存后会自动同步到系统配置，无需手动添加。')
note('删除某个选项时，若仍有画师在使用该值，系统会拒绝并列出使用该值的画师名单。')

# ══════════════════════════════════════════════════════════
# 保存
# ══════════════════════════════════════════════════════════
out_path = r'C:\Users\xxwf\IdeaProjects\supplier-portal-standalone\供应商平台使用说明.docx'
doc.save(out_path)
print(f'已生成：{out_path}')
